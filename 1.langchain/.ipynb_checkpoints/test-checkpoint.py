# Skills 封装（Runtime + Tools + 辅助函数）
# 只封装执行层与工具层；用 create_agent / create_deep_agent 挂上 skill_tools 即可。
#
# 用法示例:
#   runtime = ClaudeSkillRuntime(skills_dir="./skills", output_dir="./outputs")
#   skill_tools = make_skill_tools(runtime)
#   agent = create_deep_agent(model=llm, tools=skill_tools, system_prompt=build_skills_prompt(runtime.output_dir), ...)
#   chat_stream(agent, build_msg("..."))

from __future__ import annotations

import base64
import json
import os
import re
import shlex
import shutil
import subprocess
import sys
import time
from dataclasses import dataclass
from pathlib import Path
from typing import Any

import yaml
from langchain.messages import HumanMessage
from langchain.tools import tool

# ---------------------------------------------------------------------------
# ClaudeSkillRuntime（执行层）
# ---------------------------------------------------------------------------

@dataclass
class SkillInfo:
    name: str
    description: str
    path: Path
    skill_md: Path
    content: str


class ClaudeSkillRuntime:
    """Claude Skills runtime: SKILL.md + execute + collect outputs."""

    SCRIPT_SUFFIXES = {".py", ".js", ".mjs", ".sh", ".bat", ".ps1", ".ts"}
    OUTPUT_SUFFIXES = {".docx", ".pdf", ".pptx", ".xlsx", ".csv", ".png", ".md", ".txt", ".json", ".zip"}
    OUTPUT_IGNORE_NAMES = {"package.json", "package-lock.json"}

    def __init__(
        self,
        skills_dir: str | Path,
        work_dir: str | Path | None = None,
        output_dir: str | Path | None = None,
    ):
        self.skills_dir = Path(skills_dir).resolve()
        self.work_dir = Path(work_dir or Path.cwd()).resolve()
        self.output_dir = Path(output_dir or self.work_dir).resolve()
        self.output_dir.mkdir(parents=True, exist_ok=True)
        self.skills: dict[str, SkillInfo] = {}
        self._started_at = time.time()
        self.reload()

    def reload(self) -> None:
        self.skills.clear()
        if not self.skills_dir.is_dir():
            return
        for path in sorted(self.skills_dir.iterdir()):
            skill_md = path / "SKILL.md"
            if not path.is_dir() or not skill_md.exists():
                continue
            content = skill_md.read_text(encoding="utf-8")
            meta: dict[str, Any] = {}
            m = re.match(r"^---\s*\n(.*?)\n---\s*\n", content, re.DOTALL)
            if m:
                try:
                    meta = yaml.safe_load(m.group(1)) or {}
                except Exception:
                    meta = {}
            name = str(meta.get("name") or path.name)
            self.skills[name] = SkillInfo(
                name=name,
                description=str(meta.get("description") or "").strip(),
                path=path,
                skill_md=skill_md,
                content=content,
            )

    def list_skills(self) -> list[dict[str, str]]:
        return [
            {
                "name": s.name,
                "description": s.description,
                "skill_md": f"/skills/{s.path.name}/SKILL.md",
                "host_path": str(s.path),
            }
            for s in self.skills.values()
        ]

    def load_skill(self, name: str) -> dict[str, Any]:
        skill = self._get(name)
        return {
            "name": skill.name,
            "description": skill.description,
            "skill_md_virtual": f"/skills/{skill.path.name}/SKILL.md",
            "host_path": str(skill.path),
            "content": skill.content,
            "scripts": self.list_scripts(name),
            "hint": (
                "这是说明书，不是自动执行器。"
                "请按 content 编写代码，再用 run_js_code / run_python_code / create_docx_file 真正执行并生成文件。"
                f"生成文件请写入 OUTPUT_DIR={self.output_dir}（环境变量 OUTPUT_DIR 已注入）。"
            ),
        }

    def list_scripts(self, name: str) -> list[str]:
        skill = self._get(name)
        scripts = []
        for f in sorted(skill.path.rglob("*")):
            if not f.is_file() or f.suffix.lower() not in self.SCRIPT_SUFFIXES:
                continue
            if any(p in {"node_modules", ".git", "__pycache__"} for p in f.parts):
                continue
            scripts.append(f.relative_to(skill.path).as_posix())
        return scripts

    def run_script(self, name: str, script_relpath: str, args: str = "") -> dict[str, Any]:
        skill = self._get(name)
        script = (skill.path / script_relpath).resolve()
        if not str(script).startswith(str(skill.path.resolve())):
            return {"ok": False, "error": "脚本路径越界"}
        if not script.is_file():
            return {"ok": False, "error": f"脚本不存在: {script_relpath}"}
        arg_list = shlex.split(args, posix=os.name != "nt") if args.strip() else []
        cmd = self._cmd_for_script(script, arg_list)
        before = self._output_snapshot()
        before_work = self._work_dir_root_output_names()
        result = self._run(cmd, cwd=self.work_dir, extra={"skill": name, "script": script_relpath})
        result["new_files"] = self._collect_new_outputs(before, before_work)
        return result

    def run_python_code(self, code: str, filename: str = "_skill_runtime_tmp.py") -> dict[str, Any]:
        path = self._write_code(filename, code)
        before = self._output_snapshot()
        before_work = self._work_dir_root_output_names()
        result = self._run(["python", str(path)], cwd=self.work_dir, extra={"file": str(path)})
        result["new_files"] = self._collect_new_outputs(before, before_work)
        return result

    def run_js_code(self, code: str, filename: str = "_skill_runtime_tmp.js") -> dict[str, Any]:
        self.ensure_js_deps(["docx"])
        path = self._write_code(filename, code)
        before = self._output_snapshot()
        before_work = self._work_dir_root_output_names()
        result = self._run(["node", str(path)], cwd=self.work_dir, extra={"file": str(path)})
        result["new_files"] = self._collect_new_outputs(before, before_work)
        return result

    def ensure_js_deps(self, packages: list[str]) -> dict[str, Any]:
        pkg_json = self.work_dir / "package.json"
        node_modules = self.work_dir / "node_modules"
        logs = []
        npm = self._resolve_bin("npm")
        if not pkg_json.exists():
            r = subprocess.run([npm, "init", "-y"], cwd=str(self.work_dir), capture_output=True, text=True)
            logs.append({"step": "npm init", "returncode": r.returncode, "stderr": (r.stderr or "")[-500:]})
        missing = [p for p in packages if not (node_modules / p).exists()]
        if missing:
            r = subprocess.run(
                [npm, "install", *missing, "--save"],
                cwd=str(self.work_dir),
                capture_output=True,
                text=True,
                timeout=300,
            )
            logs.append({
                "step": f"npm install {' '.join(missing)}",
                "returncode": r.returncode,
                "stderr": (r.stderr or "")[-800:],
            })
        return {"ok": True, "missing_installed": missing, "logs": logs}

    def list_outputs(self, since_epoch: float | None = None) -> list[str]:
        since = since_epoch if since_epoch is not None else self._started_at
        files = []
        for f in self.output_dir.rglob("*"):
            if not f.is_file() or f.name in self.OUTPUT_IGNORE_NAMES:
                continue
            if any(p in {"node_modules", ".git", "__pycache__", "skills"} for p in f.parts):
                continue
            if f.suffix.lower() not in self.OUTPUT_SUFFIXES:
                continue
            if f.stat().st_mtime >= since - 1:
                files.append(str(f.relative_to(self.output_dir)).replace("\\", "/"))
        return sorted(files)

    def create_docx(
        self,
        output_filename: str,
        title: str,
        paragraphs: list[str] | None = None,
        bullet_points: list[str] | None = None,
    ) -> dict[str, Any]:
        try:
            from docx import Document
            from docx.enum.text import WD_ALIGN_PARAGRAPH
        except ImportError:
            return {"ok": False, "error": "缺少 python-docx，请先 pip install python-docx"}

        paragraphs = paragraphs or []
        bullet_points = bullet_points or []
        out = (self.output_dir / output_filename).resolve()
        if out.suffix.lower() != ".docx":
            out = out.with_suffix(".docx")
        if not str(out).startswith(str(self.output_dir)):
            return {"ok": False, "error": "输出路径必须在 OUTPUT_DIR 内"}

        out.parent.mkdir(parents=True, exist_ok=True)
        doc = Document()
        h = doc.add_heading(title, level=1)
        h.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for p in paragraphs:
            doc.add_paragraph(str(p))
        for b in bullet_points:
            doc.add_paragraph(str(b), style="List Bullet")
        doc.save(str(out))
        rel = str(out.relative_to(self.output_dir)).replace("\\", "/")
        return {
            "ok": True,
            "file": rel,
            "abs_path": str(out),
            "output_dir": str(self.output_dir),
            "size": out.stat().st_size,
        }

    def _get(self, name: str) -> SkillInfo:
        if name not in self.skills:
            raise KeyError(f"Skill 不存在: {name}；可用: {list(self.skills)}")
        return self.skills[name]

    def _write_code(self, filename: str, code: str) -> Path:
        path = (self.work_dir / Path(filename).name).resolve()
        path.write_text(code, encoding="utf-8")
        return path

    def _cmd_for_script(self, script: Path, arg_list: list[str]) -> list[str]:
        suffix = script.suffix.lower()
        if suffix == ".py":
            return ["python", str(script), *arg_list]
        if suffix in {".js", ".mjs"}:
            return ["node", str(script), *arg_list]
        if suffix == ".ts":
            return [self._resolve_bin("npx"), "--yes", "tsx", str(script), *arg_list]
        if suffix == ".ps1":
            return ["powershell", "-ExecutionPolicy", "Bypass", "-File", str(script), *arg_list]
        if suffix == ".bat":
            return [str(script), *arg_list]
        if suffix == ".sh":
            return ["bash", str(script), *arg_list]
        raise ValueError(f"不支持的脚本类型: {suffix}")

    def _run(self, cmd: list[str], cwd: Path, extra: dict | None = None) -> dict[str, Any]:
        try:
            r = subprocess.run(
                cmd,
                cwd=str(cwd),
                capture_output=True,
                text=True,
                timeout=300,
                env={
                    **os.environ,
                    "WORK_DIR": str(self.work_dir),
                    "OUTPUT_DIR": str(self.output_dir),
                    "SKILLS_DIR": str(self.skills_dir),
                },
            )
            out = {
                "ok": r.returncode == 0,
                "cmd": cmd,
                "returncode": r.returncode,
                "stdout": (r.stdout or "")[-8000:],
                "stderr": (r.stderr or "")[-8000:],
                "cwd": str(cwd),
                "output_dir": str(self.output_dir),
            }
            if extra:
                out.update(extra)
            return out
        except subprocess.TimeoutExpired:
            return {"ok": False, "error": "执行超时(>300s)", "cmd": cmd}
        except Exception as e:
            return {"ok": False, "error": str(e), "cmd": cmd}

    def _output_snapshot(self) -> set[str]:
        return set(self.list_outputs(since_epoch=0))

    def _work_dir_root_output_names(self) -> set[str]:
        if self.output_dir.resolve() == self.work_dir.resolve():
            return set()
        return {
            f.name
            for f in self.work_dir.iterdir()
            if f.is_file()
            and f.suffix.lower() in self.OUTPUT_SUFFIXES
            and f.name not in self.OUTPUT_IGNORE_NAMES
        }

    @staticmethod
    def _resolve_bin(name: str) -> str:
        # Windows: npm/npx are .cmd; subprocess without shell needs explicit path
        if sys.platform == "win32":
            found = shutil.which(f"{name}.cmd") or shutil.which(name)
            return found or f"{name}.cmd"
        return shutil.which(name) or name

    def _collect_new_outputs(self, before: set[str], before_work: set[str] | None = None) -> list[str]:
        before_work = before_work if before_work is not None else set()
        if self.output_dir.resolve() != self.work_dir.resolve():
            for name in self._work_dir_root_output_names() - before_work:
                src = self.work_dir / name
                dest = self.output_dir / name
                if dest.exists():
                    dest.unlink()
                shutil.move(str(src), str(dest))
        after = set(self.list_outputs(since_epoch=0))
        return sorted(after - before)


# ---------------------------------------------------------------------------
# make_skill_tools（@tool 层）
# ---------------------------------------------------------------------------

def make_skill_tools(runtime: ClaudeSkillRuntime) -> list:
    """Create LangChain tools bound to one Runtime instance."""

    @tool
    def list_available_skills() -> str:
        """列出所有可用 Claude Skills。"""
        return json.dumps(runtime.list_skills(), ensure_ascii=False, indent=2)

    @tool
    def load_skill(skill_name: str) -> str:
        """加载指定 Skill 的完整 SKILL.md 说明书与脚本列表。这是执行前的必做步骤。"""
        try:
            return json.dumps(runtime.load_skill(skill_name), ensure_ascii=False, indent=2)
        except Exception as e:
            return json.dumps({"ok": False, "error": str(e)}, ensure_ascii=False)

    @tool
    def list_skill_scripts(skill_name: str) -> str:
        """列出 skill 目录下可执行脚本（相对路径）。"""
        try:
            return json.dumps(
                {"skill": skill_name, "scripts": runtime.list_scripts(skill_name)},
                ensure_ascii=False,
                indent=2,
            )
        except Exception as e:
            return json.dumps({"ok": False, "error": str(e)}, ensure_ascii=False)

    @tool
    def run_skill_script(skill_name: str, script_relpath: str, script_args: str = "") -> str:
        """【Runtime 主入口】选中 skill 后必须调用本工具执行该 skill 目录内脚本并落盘。"""
        return json.dumps(
            runtime.run_script(skill_name, script_relpath, script_args),
            ensure_ascii=False,
            indent=2,
        )

    @tool
    def run_js_code(code: str, filename: str = "_skill_runtime_tmp.js") -> str:
        """将 JavaScript 写入工作目录并执行。文件应写到 process.env.OUTPUT_DIR。"""
        return json.dumps(runtime.run_js_code(code, filename), ensure_ascii=False, indent=2)

    @tool
    def run_python_code(code: str, filename: str = "_skill_runtime_tmp.py") -> str:
        """将 Python 写入工作目录并执行。生成文件请写到 os.environ['OUTPUT_DIR']。"""
        return json.dumps(runtime.run_python_code(code, filename), ensure_ascii=False, indent=2)

    @tool
    def create_docx_file(
        output_filename: str,
        title: str,
        paragraphs_json: str = "[]",
        bullet_points_json: str = "[]",
    ) -> str:
        """docx 可靠落盘工具：用 python-docx 生成 .docx（写入 OUTPUT_DIR）。"""
        try:
            paragraphs = json.loads(paragraphs_json) if paragraphs_json else []
            bullets = json.loads(bullet_points_json) if bullet_points_json else []
            if not isinstance(paragraphs, list):
                paragraphs = [str(paragraphs)]
            if not isinstance(bullets, list):
                bullets = [str(bullets)]
        except json.JSONDecodeError as e:
            return json.dumps({"ok": False, "error": f"JSON 解析失败: {e}"}, ensure_ascii=False)
        return json.dumps(
            runtime.create_docx(output_filename, title, paragraphs, bullets),
            ensure_ascii=False,
            indent=2,
        )

    @tool
    def list_generated_files() -> str:
        """列出 OUTPUT_DIR 中新生成的产物。任务结束前必须调用以确认文件已生成。"""
        return json.dumps(
            {
                "files": runtime.list_outputs(),
                "output_dir": str(runtime.output_dir),
                "work_dir": str(runtime.work_dir),
            },
            ensure_ascii=False,
            indent=2,
        )

    return [
        list_available_skills,
        load_skill,
        list_skill_scripts,
        run_skill_script,
        run_js_code,
        run_python_code,
        create_docx_file,
        list_generated_files,
    ]


# ---------------------------------------------------------------------------
# 辅助函数
# ---------------------------------------------------------------------------

def build_skills_prompt(output_dir: str | Path) -> str:
    """生成 skills 系统提示，注入当前 OUTPUT_DIR。"""
    return f"""你是 skills_agent：负责「选 skill + 调 Runtime 执行」。

## 固定两步（缺一不可）
1. **判断**：list_available_skills → load_skill(skill_name)
2. **执行（必须 tool call，禁止只口头说「我将执行」）**：
   - skill 目录已有脚本 → 立刻 run_skill_script(skill_name, script_relpath)
   - 无现成脚本、要生成 Word → create_docx_file 或 run_js_code
   - 其他 → run_python_code / run_js_code
3. **验收**：list_generated_files；无产物则根据 stderr 修复再执行

## 禁止
- 用 read_file / execute 通读或旁路执行
- Windows 绝对路径；只用相对路径或 skill 相对脚本路径
- 只描述不调用 Runtime 工具

## 约束
- 生成文件必须写入 OUTPUT_DIR（当前：{output_dir}），不要写进 skills/
- 用 run_js_code / run_python_code 时，通过 process.env.OUTPUT_DIR / os.environ['OUTPUT_DIR'] 取路径
- 最终回复给出文件相对 OUTPUT_DIR 的路径与大小
"""


def encode_image(path: str | Path) -> str:
    with open(path, "rb") as f:
        return base64.b64encode(f.read()).decode()


def build_msg(text: str | None = None, image_path: str | Path | None = None) -> list:
    """构造多模态 HumanMessage 列表。"""
    content = []
    if text:
        content.append({"type": "text", "text": text})
    if image_path:
        content.append({
            "type": "image_url",
            "image_url": {
                "url": f"data:image/jpeg;base64,{encode_image(image_path)}",
                "detail": "auto",
            },
        })
    return [HumanMessage(content=content)]


def chat_stream(agent, message, config=None):
    """流式打印 agent 执行过程。"""
    result = agent.stream({"messages": message}, stream_mode="updates", config=config)
    for chunk in result:
        for _node_name, update in chunk.items():
            if not update or "messages" not in update:
                continue
            for msg in update["messages"]:
                if msg.type == "system":
                    print(f"[{msg.type}] {msg.content}")
                if msg.type == "human":
                    print(f"[{msg.type}] {msg.content}")
                if msg.type == "ai":
                    if msg.content:
                        print(f"[{msg.type}] {msg.content}")
                    if getattr(msg, "tool_calls", None):
                        print(f"请求调用: {[tc['name'] for tc in msg.tool_calls]}")
                if msg.type == "tool":
                    print(f"[调用工具 {msg.name}]")
                    print(f"[{msg.type}] {msg.content}")


# ---------------------------------------------------------------------------
# 使用示例
# ---------------------------------------------------------------------------

def main() -> None:
    from langchain_openai import ChatOpenAI
    from deepagents import create_deep_agent
    from deepagents.backends import LocalShellBackend
    from deepagents.middleware._tool_exclusion import _ToolExclusionMiddleware

    # 脚本所在目录（等同于 notebook 的 cwd）
    current_dir = Path(__file__).resolve().parent
    SKILLS_DIR = current_dir / "skills"
    OUTPUT_DIR = current_dir / "outputs"

    llm = ChatOpenAI(
        model="qwen3.5-plus",
        openai_api_key="sk-bvOWJIfQ8Y8IrZhzzpQZ80zSKHjWIHCkPXEgQ0H4Li5UCpUC",
        openai_api_base="http://ai.wenmodel.com/v1",
        temperature=0.7,
        max_tokens=4096,
    )

    runtime = ClaudeSkillRuntime(
        skills_dir=SKILLS_DIR,
        work_dir=current_dir,
        output_dir=OUTPUT_DIR,
    )
    skill_tools = make_skill_tools(runtime)

    print("skills:", [s["name"] for s in runtime.list_skills()])
    print("output_dir:", runtime.output_dir)
    print("tools:", [t.name for t in skill_tools])

    backend = LocalShellBackend(
        root_dir=str(current_dir),
        virtual_mode=True,
        inherit_env=True,
        timeout=300,
    )

    # 屏蔽 deep agent 内置文件工具，避免抢 Runtime 执行权
    EXCLUDE = frozenset({"read_file", "write_file", "edit_file", "execute"})

    skills_agent = create_deep_agent(
        model=llm,
        backend=backend,
        tools=skill_tools,
        skills=["/skills/"],
        system_prompt=build_skills_prompt(runtime.output_dir),
        middleware=[_ToolExclusionMiddleware(excluded=EXCLUDE)],
        name="skills-agent",
    )

    print("skills_agent ready")

    msg = build_msg(text="你有哪些skills，请查询后回答")
    chat_stream(skills_agent, msg)

    msg = build_msg(text="将skills做成一个ppt文档，并保存在outputs目录下，文件名称为skills.pptx")
    chat_stream(skills_agent, msg)



if __name__ == "__main__":
    main()
